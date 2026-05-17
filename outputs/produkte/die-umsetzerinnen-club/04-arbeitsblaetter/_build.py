"""Build Arbeitsblatt-Dokumente für Die Umsetzerinnen Onboarding + Q1-Call 1."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

PETROL = RGBColor(0x12, 0x82, 0x8C)
DUNKEL = RGBColor(0x29, 0x55, 0x6D)
BODY = RGBColor(0x22, 0x22, 0x22)
OUT = Path(__file__).parent


def style_doc(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Source Sans 3"
    style.font.size = Pt(11)
    style.font.color.rgb = BODY
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = PETROL
    run.font.name = "Philosopher"
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = DUNKEL
    run.font.name = "Philosopher"
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)


def lead(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(10)


def body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def task(doc: Document, label: str, lines: int = 3) -> None:
    h = doc.add_paragraph()
    run = h.add_run(f"▸ {label}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = DUNKEL
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)
    for _ in range(lines):
        p = doc.add_paragraph("_______________________________________________________________________")
        p.paragraph_format.space_after = Pt(2)


def checkbox(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(f"☐  {it}")
        p.paragraph_format.space_after = Pt(3)


def divider(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run("─" * 65)
    run.font.color.rgb = PETROL
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)


# ─────────────────────────────────────────────────────────
# DOC 1 — Dein 30-Tage-Sprint
# ─────────────────────────────────────────────────────────
def build_sprint() -> None:
    doc = Document()
    style_doc(doc)

    h1(doc, "Dein 30-Tage-Sprint")
    lead(doc,
         "Wenn du dieses Arbeitsblatt heute ausfüllst, hast du etwas, "
         "das die meisten Mama-Mompreneurs nie haben: einen klaren, "
         "messbaren Plan für die nächsten 30 Tage. Nimm dir 15 Minuten. "
         "Schreib in Stichpunkten, nicht in Sätzen. Fertig besser als perfekt.")

    divider(doc)

    h2(doc, "1 · Wo stehe ich heute?")
    body(doc, "Ein ehrlicher Satz — was läuft, was hängt?")
    task(doc, "Was läuft gerade gut in meinem Business?", 2)
    task(doc, "Was hängt seit Wochen / Monaten?", 2)

    h2(doc, "2 · Mein 30-Tage-Sprint (EIN Ziel)")
    body(doc, "Die 4 Kriterien zum Abhaken:")
    checkbox(doc, [
        "Konkret (nicht „mehr posten“ — sondern „3 Karussells veröffentlicht“)",
        "Messbar (Zahl drin oder klare Ja/Nein-Antwort am Tag 30)",
        "In 30 Tagen erreichbar (kein 6-Monats-Plan)",
        "Bringt mein Business wirklich voran (nicht nur den Kalender)",
    ])
    task(doc, "MEIN 30-TAGE-SPRINT (in einem Satz):", 3)

    h2(doc, "3 · Warum genau dieser Sprint?")
    body(doc, "Wenn du das in einem Satz beantworten kannst, hältst du dran fest, "
              "auch wenn der Mama-Alltag mal wieder zuschlägt.")
    task(doc, "Warum genau dieses Ziel — und warum jetzt?", 3)

    h2(doc, "4 · Die ehrliche Hürden-Liste")
    body(doc, "Was wird dich am meisten ablenken? Schreib es VORHER auf — "
              "dann erkennst du es, wenn es kommt.")
    task(doc, "Top 3 Hürden / Ablenkungen, die ich diesen Monat erwarten kann:", 3)

    h2(doc, "5 · Mein Wochen-Schritt")
    body(doc, "Teile den Sprint in 4 Wochen. Pro Woche: 1 konkreter, kleiner Schritt.")
    task(doc, "Woche 1 — bis [Datum]:", 1)
    task(doc, "Woche 2 — bis [Datum]:", 1)
    task(doc, "Woche 3 — bis [Datum]:", 1)
    task(doc, "Woche 4 — bis [Datum]:", 1)

    h2(doc, "6 · Mein wöchentliches Review (5 Min, jeden Sonntag)")
    body(doc, "Antworte einmal pro Woche auf diese 3 Fragen — nicht mehr.")
    checkbox(doc, [
        "Habe ich diese Woche meinen Schritt gemacht? (Ja / Nein / Halb)",
        "Was hat mich aufgehalten?",
        "Was ist mein konkreter nächster Schritt?",
    ])

    h2(doc, "7 · Sichtbar machen")
    body(doc, "Sichtbares Commitment hält dran. Wähle EINEN Weg:")
    checkbox(doc, [
        "Ich teile meinen Sprint im Telegram-Kanal mit #Sprint30",
        "Ich schreibe ihn auf einen Zettel und klebe ihn an den Spiegel",
        "Ich erzähle ihn mindestens 2 Menschen aus meinem Umfeld",
    ])

    divider(doc)
    h2(doc, "Frage für den nächsten Call?")
    body(doc, "Wenn du Sparring zu deinem Sprint brauchst, schreib deine Frage hier auf "
              "und poste sie ins Telegram unter #FrageCall — dann ist sie sicher dran.")
    task(doc, "Meine Frage:", 3)

    out_path = OUT / "30-tage-sprint.docx"
    doc.save(out_path)
    print(f"[OK] {out_path}")


# ─────────────────────────────────────────────────────────
# DOC 2 — Q1-Funnel-Audit
# ─────────────────────────────────────────────────────────
def build_funnel_audit() -> None:
    doc = Document()
    style_doc(doc)

    h1(doc, "Funnel-Audit in 30 Minuten")
    lead(doc,
         "Begleitendes Arbeitsblatt zum Q1-Call „Funnel-Optimierung — wo dein Geld liegen bleibt“. "
         "Du brauchst nichts ausser einem Stift und 30 Minuten ehrlicher Selbstbeobachtung. "
         "Schätzen ist erlaubt. Fertig besser als perfekt.")

    divider(doc)

    h2(doc, "Schritt 1 · Mein Funnel auf einen Blick (5 Min)")
    body(doc, "Skizziere deinen aktuellen Funnel — von Sichtbarkeit bis Stammkundin. "
              "Wenn eine Stufe fehlt: kreise sie ein. Eine fehlende Stufe ist auch ein Befund.")
    for stage, hint in [
        ("SICHTBAR — wo werden Menschen zum ersten Mal auf mich aufmerksam?",
         "(Reels / Karussells / Stories / Werbeanzeigen / Empfehlungen / ?)"),
        ("INTERESSIERT — was klicken sie an, um mehr zu erfahren?",
         "(Bio-Link / Karussell-CTA / DM-Keyword / Werbe-Klick / ?)"),
        ("LEAD — wo geben sie ihre E-Mail / kontaktieren mich?",
         "(Freebie-Opt-in / DM / ManyChat-Flow / Quiz / ?)"),
        ("KÄUFERIN — wann + wie kaufen sie das erste Produkt?",
         "(direkt nach Opt-in / nach Mailsequenz / nach Webinar / ?)"),
        ("STAMMKUNDIN — was passiert nach dem ersten Kauf?",
         "(Upsell / Folge-Mail / Bundle / Membership / ?)"),
    ]:
        task(doc, stage, 1)
        body(doc, f"Beispiele: {hint}")

    h2(doc, "Schritt 2 · Meine echten Zahlen (10 Min)")
    body(doc, "Trag deine letzten 30 Tage ein. Wenn du keine Zahlen hast: schätze. "
              "Schätzen ist besser als nichts. Notiere wo du schätzt — du kannst nächstes Mal messen.")
    task(doc, "Reichweite letzter 30 Tage (ungefähr):", 1)
    task(doc, "Klicks auf Bio-Link / Funnel-Einstieg:", 1)
    task(doc, "Neue E-Mail-Leads:", 1)
    task(doc, "Verkäufe (egal welcher Höhe):", 1)
    task(doc, "Wiederkäufe / zweite Buchungen:", 1)

    h2(doc, "Schritt 3 · Vergleich mit Standard-Verlustraten (5 Min)")
    body(doc, "So sehen typische Verlustraten aus:")
    checkbox(doc, [
        "Sichtbar → Klick: ca. 10 % schaffen es",
        "Klick → Lead: ca. 40 % schaffen es",
        "Lead → Käuferin: ca. 5 % schaffen es",
        "Käuferin → Stammkundin: ca. 50 % schaffen es",
    ])
    body(doc, "Wo sind deine Zahlen WEIT unter dem Schnitt? Dort liegt dein grösstes Leck.")
    task(doc, "Mein grösster Verlust passiert zwischen Stufe ___ und Stufe ___:", 2)

    h2(doc, "Schritt 4 · Welches Leck habe ich? (5 Min)")
    body(doc, "Ordne dich einer der 4 Diagnosen zu:")
    checkbox(doc, [
        "Leck #1 — Hook spricht die Falsche an  (Reichweite okay, kaum Klicks)",
        "Leck #2 — Landingpage verkauft Inhalte statt Transformation  (Klicks da, kaum Leads)",
        "Leck #3 — Keine Mail-Sequenz nach Opt-in  (Leads wachsen, kaum Verkäufe)",
        "Leck #4 — Kein Folge-Angebot nach Erstkauf  (Käuferinnen kaufen einmal)",
    ])
    task(doc, "Mein grösstes Leck ist #___:", 1)

    h2(doc, "Schritt 5 · Mein 1-Wochen-Fix (5 Min)")
    body(doc, "Wähle EINE Sache. Nicht alles. Eine — und mach sie diese Woche.")
    task(doc, "Mein konkreter Fix für diese Woche:", 2)
    task(doc, "Bis wann genau? (Datum + Uhrzeit, sonst passiert es nicht):", 1)
    task(doc, "Wer weiss davon? (Sichtbarkeit hält dran):", 1)

    divider(doc)
    h2(doc, "Sparring im Call?")
    body(doc, "Wenn du beim Audit irgendwo hängst — bring die Frage in den Q&A-Call mit. "
              "Poste sie vorher im Telegram unter #FunnelLeck.")
    task(doc, "Meine Funnel-Frage für den Call:", 3)

    out_path = OUT / "q1-funnel-audit.docx"
    doc.save(out_path)
    print(f"[OK] {out_path}")


if __name__ == "__main__":
    build_sprint()
    build_funnel_audit()
    print("Done.")
