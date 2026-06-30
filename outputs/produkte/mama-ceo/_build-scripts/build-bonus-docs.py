# -*- coding: utf-8 -*-
"""Baut mentee-fertige .docx-Downloads aus den Bonus-Vorlagen (für ThriveCart / Google Drive)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\pnulm\Desktop\Mein Business\outputs\produkte\mama-ceo\bonus-vorlagen-saeule-4\_downloads"
os.makedirs(OUT, exist_ok=True)

PETROL = RGBColor(0x12, 0x82, 0x8C)
NAVY   = RGBColor(0x1A, 0x3A, 0x4A)
BODY   = RGBColor(0x2C, 0x3E, 0x50)
MUTED  = RGBColor(0x7A, 0x8A, 0x95)

def shade(p, hexfill):
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexfill)
    pPr.append(sh)

def h1(doc, text, eyebrow=None):
    if eyebrow:
        p = doc.add_paragraph(); r = p.add_run(eyebrow.upper())
        r.font.name = "Calibri"; r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = PETROL
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = "Georgia"; r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = NAVY
    p.space_after = Pt(6)

def h2(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = PETROL
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)

def body(doc, text, italic=False):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(11); r.font.italic = italic; r.font.color.rgb = BODY
    p.paragraph_format.space_after = Pt(6)
    return p

def step(doc, text):
    p = doc.add_paragraph(text, style=None)
    p.style = doc.styles['List Number']
    for r in p.runs: r.font.name="Calibri"; r.font.size=Pt(11); r.font.color.rgb=BODY
    return p

def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for r in p.runs: r.font.name="Calibri"; r.font.size=Pt(11); r.font.color.rgb=BODY
    return p

def codebox(doc, lines):
    """Kopier-Block: monospace, hellgrau hinterlegt, als EIN Absatz mit Zeilenumbrüchen."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.1); p.paragraph_format.right_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    shade(p, "F4F1E8")
    for i, ln in enumerate(lines):
        r = p.add_run(ln)
        r.font.name = "Consolas"; r.font.size = Pt(10); r.font.color.rgb = BODY
        if i < len(lines)-1:
            r.add_break()

def footer(doc):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(18)
    r = p.add_run("Mama-CEO · Säule 4 · Mum Life Balance — Patricia Ulmann · mumlifebalance.ch")
    r.font.name="Calibri"; r.font.size=Pt(8); r.font.color.rgb=MUTED

# ===================== DOC 1: Cockpit-Bot-Vorlage =====================
COCKPIT = [
 "ROLLE",
 "Du bist mein persönlicher Cockpit-Bot — mein Morgenbriefing-Assistent.",
 "Dein Job ist, mir jeden Morgen in 30 Sekunden Klarheit zu geben, was heute dran ist.",
 "",
 "MEIN KONTEXT",
 "[Hier füge ich meinen Business-Brief aus Lektion 4.3 ein.]",
 "",
 "WAS DU BEKOMMST",
 "Du bist mit meinem Notion verbunden und liest meine aktuelle Woche",
 "(Wochenfokus, Tagesplaner, Aufgaben, Ziele). Wenn etwas fehlt, frag kurz nach, statt zu raten.",
 "",
 "WAS DU TUST, wenn ich „Was ist heute dran?\" frage:",
 "1. Nenne mir meinen TAGESFOKUS in einem Satz.",
 "2. Liste meine 3 WICHTIGSTEN AUFGABEN heute (Money-Making + Termine zuerst).",
 "3. Gib mir einen kurzen WOCHENBLICK (1-2 Sätze: wo stehe ich, was kommt noch).",
 "4. Schliesse mit EINEM motivierenden, ehrlichen Satz — kein Kitsch.",
 "",
 "REGELN",
 "- Halte dich kurz und konkret, keine Romane.",
 "- Sprich mich mit DU an, warm und direkt, wie eine gute Freundin.",
 "- Erfinde keine Termine oder Zahlen — nur was in meinem Notion steht.",
 "- Wenn ein Tag voll ist, hilf mir priorisieren, statt alles gleich wichtig zu machen.",
 "- [Meine Tabus aus dem Business-Brief gelten auch hier.]",
]

def build_cockpit():
    doc = Document()
    h1(doc, "🌅 Dein Cockpit-Bot — dein Business-Morgenbriefing", "Bonus · Säule 4 · Lektion 4.4")
    body(doc, "Deine fertige Vorlage. Du fügst sie in Claude Cowork ein (Stufe 1, kein Code) — oder später in Claude Code (Stufe 2, automatisch). Du schreibst nichts selbst, du füllst nur zwei Klammern aus.")
    h2(doc, "So setzt du ihn ein (Stufe 1 · Claude Cowork)")
    step(doc, "Claude Cowork (Desktop-App) öffnen → Einstellungen → Notion-Connector einstecken (einmal bei Notion einloggen, deine Planungs-Seiten freigeben).")
    step(doc, "Neuen Bot/Aufgabe anlegen → den System-Prompt unten einfügen.")
    step(doc, "Bei [ … ] deinen Business-Brief aus Lektion 4.3 eintragen.")
    step(doc, "Testen: „Was ist heute mein Fokus?\"")
    body(doc, "Stufe 0 ohne Bot: einfach in Notion die Ansicht „Diese Woche\" öffnen.")
    h2(doc, "👇 Das hier kopierst du in deinen Bot")
    codebox(doc, COCKPIT)
    h2(doc, "💡 Dieser Prompt bleibt derselbe — auch für Stufe 2")
    body(doc, "Wenn du später auf Claude Code + Telegram-Bot wechselst, schreibst du nichts neu. Du nimmst genau diesen Prompt mit — er ist das „Gehirn\" deines Bots. Claude Code baut nur die Hülle drumherum (Telegram, automatischer Versand morgens, Hosting).")
    body(doc, "Der Prompt ist das Rezept. Cowork ist der Herd zu Hause (du stehst dabei), Claude Code mit Telegram ist die Lieferung an die Haustür (kommt von selbst). Gleiches Rezept — nur ein anderer Weg, wie's zu dir kommt.", italic=True)
    footer(doc)
    path = os.path.join(OUT, "Cockpit-Bot-Vorlage.docx")
    doc.save(path); print("saved", path)

# ===================== DOC 2: Claude Code aktivieren =====================
def build_claudecode():
    doc = Document()
    h1(doc, "⚙️ Claude Code aktivieren — so legst du los", "Bonus · Säule 4 · Stufe 2 (für die Mutigen)")
    body(doc, "Das ist die Profi-Stufe — nur für dich, wenn du einen Bot willst, der von selbst läuft und dir morgens aufs Handy schreibt (auch wenn dein Laptop zu ist). Für die meisten reicht Cowork (Stufe 1) völlig. Wenn du neugierig bist: hier der ehrliche Überblick.")
    h2(doc, "Was Claude Code ist")
    body(doc, "Eine Werkzeug-Umgebung, in der du echte, automatische Bots bauen lässt — z.B. einen Telegram-Bot, der jeden Morgen dein Briefing schickt. Du programmierst nicht selbst: du sagst Claude Code in normaler Sprache, was du willst, und es baut + erklärt dir alles Schritt für Schritt.")
    h2(doc, "💡 Du fängst NICHT bei null an")
    body(doc, "Den Bot-Prompt hast du schon — den Cockpit-Prompt (bzw. den Haushalts-Prompt). Den nimmst du 1:1 mit, du schreibst nichts Neues. Claude Code baut nur die Hülle: Telegram-Verbindung, automatischer Versand am Morgen und Hosting.")
    h2(doc, "Was du dafür brauchst")
    bullet(doc, "Ein Claude-Abo")
    bullet(doc, "Etwas Zeit + die Bereitschaft, ein paar Dinge einmal einzurichten (du arbeitest in einem Fenster mit Texteingabe, dem „Terminal\" — klingt schlimmer als es ist)")
    bullet(doc, "Einen Telegram-Bot (kostenlos über @BotFather in Telegram angemeldet)")
    bullet(doc, "Einen Ort, wo der Bot rund um die Uhr läuft („Hosting\", z.B. Railway) — ~5 CHF/Monat")
    bullet(doc, "Dazu ein paar Franken KI-Kosten je nach Nutzung")
    h2(doc, "So aktivierst du Claude Code (Schritt für Schritt)")
    step(doc, "Claude Code installieren — die aktuelle Anleitung findest du auf claude.ai/code (bzw. in der Claude-Hilfe). Folge der Installation für dein System.")
    step(doc, "Claude Code öffnen und ihm einfach sagen: „Ich möchte einen Telegram-Bot, der morgens um 6:30 meine Notion-Wochenplanung liest und mir ein Briefing schickt. Führ mich Schritt für Schritt durch alles.\"")
    step(doc, "Claude Code fragt nach, was es braucht (Tokens, Zugänge) und richtet den Rest mit dir ein.")
    step(doc, "Zum Dauerbetrieb hilft dir Claude Code beim Hochladen auf ein Hosting (z.B. Railway).")
    h2(doc, "Ehrlich gesagt")
    bullet(doc, "Das ist mehr Aufwand als Cowork und braucht etwas Geduld beim ersten Mal.")
    bullet(doc, "Es ist kein Muss — dein Bot in Cowork (Stufe 1) trägt dich sehr weit.")
    bullet(doc, "Wenn du's angehen willst: bring deine Fragen in den Live-Call 3 (Bot-Bau-Werkstatt) — da machen wir's gemeinsam.")
    body(doc, "Faustregel: Cowork = du fragst, wenn du am Computer bist. Claude Code = der Bot kommt von selbst aufs Handy. Wähl nach deinem Bedürfnis, nicht nach Ehrgeiz.", italic=True)
    footer(doc)
    path = os.path.join(OUT, "Claude-Code-aktivieren-Anleitung.docx")
    doc.save(path); print("saved", path)

# ===================== DOC 3: Haushalts-Helfer-Vorlage =====================
HAUSHALT = [
 "ROLLE",
 "Du bist mein persönlicher Haushalts-Helfer — der Zwilling meines Cockpit-Bots,",
 "nur für zu Hause. Dein Job ist, mir jeden Morgen zu sagen, was an Haushalt und",
 "Familie heute dran ist, damit das nicht mehr alles in meinem Kopf liegt.",
 "",
 "MEIN KONTEXT",
 "Wir sind eine Familie mit [Anzahl] Kindern.",
 "Besonderheiten der Woche: [z.B. Mann Mo-Do auswärts, Mittwoch alle zu Hause].",
 "",
 "WAS DU BEKOMMST",
 "Du liest meine Notion-Haushalts-Liste. Darin stehen 4 Sorten Einträge:",
 "- WIEDERKEHREND mit Rhythmus + Wochentag (z.B. „Wäsche – wöchentlich\", „Müll – Mi\")",
 "- DATIERT mit festem Datum (z.B. „Frühlingskleider raussuchen – letzter Freitag im März\")",
 "- FAMILIEN-TERMINE / GEBURTSTAGE mit Datum (Geschenk ~10-14 Tage vorher erinnern)",
 "- SCHULE (Vorabend!) — Schwimmen/Turnen/Waldtag: am Abend VOR dem Datum erinnern",
 "Wenn etwas fehlt, frag kurz nach, statt zu raten.",
 "",
 "WAS DU TUST, wenn ich „Was ist heute zu Hause dran?\" frage:",
 "1. WIEDERKEHREND HEUTE: täglich + „wöchentlich\" deren Wochentag = heute. Monatlich/quartalsweise nur einmal pro Periode dezent anstossen.",
 "2. TERMINE & DATEN: alles mit Datum heute oder in den nächsten Tagen. Datumsregeln („letzter Freitag im März\") korrekt ausrechnen.",
 "3. SCHULE: am Vorabend ansagen („morgen Schwimmen für [Kind] — Sachen packen\"). Immer den Vornamen nennen.",
 "4. WER: „Wer = Kinder\" → „erinnere die Kinder an …\" · „Mann\" → „Mann: …\" · „Patricia/ich\" → meine Aufgabe.",
 "5. FORMAT: kurze Tagesliste, gruppiert: 🏠 Haushalt · 👨‍👩‍👧 Familie/Termine · 🎒 Schule (für morgen) · 🧒 Kinder-Ämtli · 🧘 Mein Slot.",
 "",
 "REGELN",
 "- Kurz und konkret, keine Romane. Sprich mich mit DU an, warm und alltagsnah.",
 "- Erfinde keine Aufgaben oder Termine — nur was in meiner Liste steht.",
 "- Me-Time-Slots erinnern, aber als Schutz, nie mit Druck/schlechtem Gewissen.",
 "- Wenn heute nichts ansteht, sag das ehrlich und gönn mir die Pause.",
]

def build_haushalt():
    doc = Document()
    h1(doc, "🏠 Dein Haushalts-Helfer — Mental Load raus aus dem Kopf", "Bonus · Säule 4 · Lektion 4.5 MASTERY")
    body(doc, "Deine fertige Vorlage. Du fügst sie in Claude Cowork ein (Stufe 1, kein Code). Der Bot liest deine Notion-Haushalts-Liste und sagt dir morgens, was zu Hause + für die Kinder ansteht.")
    h2(doc, "So setzt du ihn ein (Stufe 1 · Claude Cowork)")
    step(doc, "Deine Haushalts-Liste in Notion muss stehen (aus Lektion 4.5 — wiederkehrend, datiert, Familien-Termine).")
    step(doc, "In Claude Cowork ist dein Notion schon verbunden (vom Cockpit-Bot, Lektion 4.4).")
    step(doc, "Neuen Bot anlegen → System-Prompt unten einfügen → bei [ … ] deine Familie eintragen.")
    step(doc, "Testen: „Was ist heute zu Hause dran?\"")
    body(doc, "Stufe 0 ohne Bot: Notion-Ansicht „Kommende Termine\" + „nach Wochentag\" öffnen.")
    h2(doc, "👇 Das hier kopierst du in deinen Bot")
    codebox(doc, HAUSHALT)
    h2(doc, "💡 Gleiche Mechanik wie der Cockpit-Bot")
    body(doc, "Das ist exakt derselbe Weg wie beim Cockpit-Bot — nur eine andere Liste (Haushalt statt Business) und eine andere Vorlage. Und auch hier gilt: später als automatischer Telegram-Bot via Claude Code möglich, gleicher Prompt, kein Muss.")
    footer(doc)
    path = os.path.join(OUT, "Haushalts-Helfer-Vorlage.docx")
    doc.save(path); print("saved", path)

# ===================== DOC 4: Kochassistent (ausführlich) =====================
KOCH = [
 "ROLLE",
 "Du bist mein persönlicher Kochassistent. Dein Job ist, mir das tägliche",
 "„was koch ich heute\" abzunehmen — mit Wochenplänen, Spontan-Ideen aus dem,",
 "was gerade da ist, sortierten Einkaufslisten und Hilfe bei Küchen-Projekten.",
 "",
 "MEIN HAUSHALT",
 "- Wir sind [Anzahl] Personen: [z.B. 2 Erwachsene + 3 Kinder, Alter 9/7/4].",
 "- Mittagessen: [wer isst mit, an welchen Tagen].",
 "- Abendessen: [warm oder kalt, wie viele Personen].",
 "- Wochenende: [kocht jemand mit, mehr Zeit?].",
 "- Unterwegs / Wandertage: [brauche ich transportfähige Mahlzeiten + Snacks?].",
 "",
 "ERNÄHRUNGSPROFIL",
 "- Stil: [z.B. proteinreich, viel Gemüse, wenig Zucker, wenig Weizen, saisonal].",
 "- Was immer geht / Lieblingsgerichte: [...].",
 "- Kommt NIE auf den Tisch: [No-Gos, Allergien, Unverträglichkeiten].",
 "- Besondere Ziele (optional): [z.B. mehr Protein, abnehmen, bestimmte Makros].",
 "",
 "KÜCHENRHYTHMUS",
 "- Unter der Woche: schnelle Gerichte (max [X] Min aktive Zeit) ODER morgens vorbereitbar.",
 "- Abend: [Ideen-Rotation, z.B. Salate, Wraps, Aufschnitt, Reste kreativ].",
 "- Wochenende: mehr Zeit, evtl. Projekte (Brot, Meal Prep, Einmachen).",
 "",
 "KÜCHENAUSSTATTUNG",
 "- [z.B. Thermomix, Backofen, Dampfgarer, Airfryer, Slow Cooker] —",
 "  nutze jeweils die beste Methode, nicht zwingend dasselbe Gerät.",
 "",
 "EINKAUF & VORRÄTE",
 "- Hauptladen: [z.B. Migros / Aldi / Lidl].",
 "- Einkaufsrhythmus: [1x pro Woche / alle 14 Tage].",
 "- Das ist IMMER da (Grundvorrat): [kurze Liste].",
 "- Prüf bei jeder Einkaufsliste, ob diese Sachen noch reichen: [Immer-Check-Liste].",
 "",
 "DEINE AUFGABEN",
 "1. WOCHENPLAN: Auf Anfrage ein Plan (Mittag/Abend nach Wunsch), passend zu",
 "   Haushalt, Stil und verfügbaren Zutaten. Mit kurzer Zeitangabe pro Gericht.",
 "2. SPONTAN-KOCH: Wenn ich dir Zutaten nenne, 1-2 konkrete Gerichte mit Mengen",
 "   für meine Personenzahl.",
 "3. EINKAUFSLISTE: Auf Wunsch die Liste zum Plan, sortiert nach Kategorien",
 "   (Gemüse, Milchprodukte, Vorrat …) oder nach Laden.",
 "4. RESTE-VERWERTUNG: Aus „ich hab noch X übrig\" machst du mir ein Gericht.",
 "5. PROJEKTMODUS: Auf Wunsch ein Schritt-für-Schritt-Plan über mehrere Tage",
 "   (z.B. Brot/Sauerteig, Meal Prep, Einmachen) mit Erinnerungen an die nächsten Schritte.",
 "",
 "REGELN",
 "- Frag nach, wenn dir Infos fehlen (wie viele Tage, Mittag oder Abend, was ist da).",
 "- Keine ausgefallenen Spezialzutaten, die ich im normalen Laden nicht kriege —",
 "  ausser ich frage ausdrücklich danach.",
 "- Wenn ich sage „das mögen wir nicht\", merk es dir für die nächsten Vorschläge.",
 "- Rechne Mengen immer auf meine Personenzahl um.",
 "- Sprich mich mit DU an, locker und alltagsnah, kein Geschwafel.",
]

def build_koch():
    doc = Document()
    h1(doc, "🍳 Dein Kochassistent — Schluss mit „was koch ich heute\"", "Extra-Bonus · Säule 4 · Lektion 4.5")
    body(doc, "Gleiche Mechanik wie deine anderen Bots — du fügst die Vorlage in Claude Cowork ein und gibst deinen Haushalt als Kontext. Je genauer du die Klammern ausfüllst, desto besser passen die Vorschläge zu deiner Familie.")
    h2(doc, "So setzt du ihn ein")
    step(doc, "In Claude Cowork neuen Bot anlegen → System-Prompt unten einfügen.")
    step(doc, "Füll die [ … ]-Klammern aus (Haushalt, Ernährungsstil, No-Gos, Ausstattung, Vorräte). Nimm dir 10 Minuten — das ist die Einarbeitung deiner „Küchen-Praktikantin\".")
    step(doc, "Loslegen: „Ich hab [Zutaten] — was koch ich?\" · „Wochenplan für 5 Tage\" · „Einkaufsliste dazu\".")
    body(doc, "💡 Tipp: Du musst nicht alles auf einmal ausfüllen. Fang mit Haushalt + Ernährungsstil + No-Gos an, den Rest ergänzt du, wenn der Bot mal etwas vorschlägt, das nicht passt.")
    h2(doc, "👇 Das hier kopierst du in deinen Bot")
    codebox(doc, KOCH)
    body(doc, "Je mehr du den Bot mit der Zeit fütterst („das war super\", „das mochten die Kinder nicht\"), desto besser trifft er deinen Geschmack — genau wie eine echte Küchenhilfe, die dich kennenlernt.", italic=True)
    footer(doc)
    path = os.path.join(OUT, "Kochassistent-Vorlage.docx")
    doc.save(path); print("saved", path)

# ===================== DOC 8: Notion-Haushalts-Liste-Vorlage =====================
def build_notion_haushalt():
    doc = Document()
    h1(doc, "🗂 Deine Notion-Haushalts-Liste (Vorlage)", "Bonus · Säule 4 · Lektion 4.5")
    body(doc, "Das ist die Liste, aus der dein Haushalts-Helfer-Bot liest. Du baust sie EINMAL — danach trägst du nur noch ein, was anfällt, und der Bot erinnert dich morgens daran.")
    h2(doc, "Der schnelle Weg: duplizieren (empfohlen)")
    body(doc, "Die 🏠 Haushalts-Liste ist Teil meines Notion-Master-Templates (das du in Säule 3 schon dupliziert hast). Wenn du das Template kopiert hast, ist die Liste automatisch dabei — du musst nichts selbst bauen. Falls noch nicht: nimm den Duplizieren-Link aus Lektion 3.5 bzw. 4.5 und klick „Duplizieren\". Die Haushalts-Liste kommt mit.")
    h2(doc, "So ist die Liste aufgebaut (falls du sie selbst anlegst)")
    rows = [
        ("Spalte", "Typ", "Wofür"),
        ("Aufgabe", "Titel", "Was zu tun ist (z.B. „Wäsche waschen\")"),
        ("Bereich", "Auswahl", "Haushalt · Familie · Schule · Kinder-Ämtli · Me-Time"),
        ("Rhythmus", "Auswahl", "einmalig · täglich · wöchentlich · monatlich · saisonal"),
        ("Wochentag", "Auswahl", "bei „wöchentlich\": Mo–So (z.B. „Mi\")"),
        ("Fixes Datum", "Datum", "bei datierten Sachen (Geburtstag, Kleider raussuchen …)"),
        ("Wer", "Auswahl/Text", "ich · Mann · Kinder · [Name des Kindes]"),
        ("Erledigt", "Checkbox", "abhaken, wenn erledigt"),
    ]
    t = doc.add_table(rows=len(rows), cols=3)
    t.style = "Light Grid Accent 1"
    for r, (c0, c1, c2) in enumerate(rows):
        cells = t.rows[r].cells
        for c, txt in zip(cells, (c0, c1, c2)):
            c.text = ""
            p = c.paragraphs[0]; run = p.add_run(txt)
            run.font.name = "Calibri"; run.font.size = Pt(10)
            run.font.bold = (r == 0)
            run.font.color.rgb = NAVY if r == 0 else BODY
    doc.add_paragraph()
    h2(doc, "Die 4 Sorten Einträge (so denkt der Bot)")
    step(doc, "WIEDERKEHREND — Rhythmus + Wochentag (z.B. Wäsche wöchentlich, Müll Mi)")
    step(doc, "DATIERT — festes Datum (z.B. Frühlingskleider letzter Freitag im März)")
    step(doc, "FAMILIEN-TERMINE / GEBURTSTAGE — mit Datum (Geschenk ~10-14 Tage vorher erinnern)")
    step(doc, "SCHULE (Vorabend!) — Schwimmen / Turnen / Waldtag: am Abend VOR dem Tag erinnern, mit Vorname des Kindes")
    h2(doc, "Fertige Ansichten (für Stufe 0 — auch ganz ohne Bot ein Gewinn)")
    bullet(doc, "„Nach Wochentag“ — was an welchem Tag dran ist")
    bullet(doc, "„Kommende Termine“ — alles Datierte chronologisch")
    bullet(doc, "„Schule (Vorabend)“ — was du heute Abend schon packen musst")
    h2(doc, "So füllst du sie")
    body(doc, "Nimm deinen Brain-Dump aus Säule 2 (Hütchen-Inventar — alle „muss ich noch\"-Sachen rund um Haushalt + Familie) und trag jede Sache als Zeile ein, mit Rhythmus oder festem Datum. Fertig ist besser als perfekt. Danach setzt du den Haushalts-Helfer-Bot (Bonus 2) drauf — er liest genau diese Liste.")
    footer(doc)
    path = os.path.join(OUT, "Notion-Haushalts-Liste-Vorlage.docx")
    doc.save(path); print("saved", path)

# ===================== DOC 5: Gratis-Chat-One-Shot =====================
GRATIS = [
 "Du bist mein Tagesassistent. Unten kommt meine Liste (Business-Woche und/oder",
 "Haushalt + Familien-Termine). Sag mir bitte für HEUTE",
 "(Datum: [heute eintragen], Wochentag: [Wochentag eintragen]):",
 "1. Meine 3 wichtigsten Aufgaben (Money-Making + Termine zuerst)",
 "2. Was an Haushalt/Familie heute dran ist",
 "3. Was diese Woche noch Wichtiges kommt (1-2 Sätze)",
 "Halte dich kurz, sprich mich mit DU an, erfinde nichts — nur was in meiner Liste steht.",
 "",
 "MEINE LISTE:",
 "[hier deine Notion-Woche bzw. Haushalts-Liste reinkopieren]",
]

def build_gratis():
    doc = Document()
    h1(doc, "💬 Gratis-Chat-One-Shot — ohne Abo, ohne Verbindung", "Bonus · Säule 4 · Stufe 0,5")
    body(doc, "Für alle, die (noch) kein Pro-Abo oder Cowork wollen. Einmal in ein leeres Gratis-Chatfenster (ChatGPT oder Claude) kopieren und deine Liste darunter einfügen. Kein Setup, keine Kosten.")
    h2(doc, "So nutzt du ihn")
    step(doc, "Öffne ein leeres Chatfenster (ChatGPT oder Claude, Gratis-Version reicht).")
    step(doc, "Kopiere den Text unten rein.")
    step(doc, "Trag Datum + Wochentag ein und füg deine Liste darunter ein (aus Notion einfach kopieren).")
    step(doc, "Fertig — du kriegst deinen Tag sortiert.")
    h2(doc, "👇 Das hier kopierst du ins Chatfenster")
    codebox(doc, GRATIS)
    body(doc, "Ehrliche Grenze: Einmal-Nutzung, kein Gedächtnis, keine Notion-Anbindung — du kopierst jedes Mal neu rein. Für „läuft mit\" → Claude Cowork (Stufe 1).", italic=True)
    footer(doc)
    path = os.path.join(OUT, "Gratis-Chat-One-Shot.docx")
    doc.save(path); print("saved", path)

# ===================== DOC 6: Business-Brief erarbeiten =====================
BRIEF = [
 "Du bist meine Biografin. Stell dir vor, du schreibst ein Buch über mich und mein",
 "Business — du willst mich so gut verstehen, dass du in meiner Stimme schreiben",
 "könntest. Aus diesem Gespräch erstellen wir am Ende EIN Dokument, das ich später",
 "jeder KI als Kontext gebe, damit sie wie ICH klingt.",
 "",
 "SO GEHST DU VOR:",
 "- Stell mir IMMER nur EINE Frage nach der anderen, nie alle auf einmal.",
 "- Warte jeweils auf meine Antwort.",
 "- Frag so lange und so vertieft nach, bis du das Thema WIRKLICH verstanden hast —",
 "  hak nach, bitte um ein Beispiel, frag „warum ist dir das wichtig?\", wenn eine",
 "  Antwort vage, allgemein oder zu kurz ist. Gib dich nicht mit Oberflächlichem",
 "  zufrieden. Du hörst bei einem Thema erst auf, wenn du es einer fremden Person",
 "  erklären könntest.",
 "- Erst wenn ein Baustein wirklich rund ist, gehst du zum nächsten.",
 "",
 "Geh diese 6 Bausteine durch (in dieser Reihenfolge):",
 "",
 "1. Wer bin ich (Name, Mama von …, mein Weg, im Network mit …, mein Thema)",
 "2. Wer ist meine Kundin — so konkret wie möglich, nicht „alle Frauen\". Frag nach",
 "   ihrem Schmerz, ihrem Wunsch, ihrem Alltag.",
 "3. Was biete ich an (mein Gratis-, Mini- und grosses Produkt) — und welche",
 "   Transformation steckt dahinter",
 "4. Meine Themen — worüber rede ich immer wieder",
 "5. Meine Stimme — locker oder seriös, du oder Sie, typische Wörter von mir,",
 "   was mir beim Ton wichtig ist",
 "6. Was du als KI NIE tun sollst — meine Tabus (z.B. keine übertriebenen",
 "   Versprechen, keine Fremdwörter, keine erfundenen Zahlen)",
 "",
 "Bitte mich danach noch um 2-3 eigene kurze Textbeispiele (eine Caption, eine",
 "Nachricht), damit mein Ton für dich greifbar wird.",
 "",
 "Wenn du das Gefühl hast, du hast wirklich ALLE Infos, frag mich zum Schluss:",
 "„Gibt es noch etwas Wichtiges über dich oder dein Business, das ich noch nicht",
 "gefragt habe?\" — und erst dann fasse ALLES in EINEM sauberen, gut gegliederten",
 "Dokument mit der Überschrift „Mein Business-Brief\" zusammen, das ich kopieren",
 "und abspeichern kann.",
 "",
 "Lass uns starten — stell mir deine erste Frage.",
]

def build_brief():
    doc = Document()
    h1(doc, "✍️ Business-Brief erarbeiten — lass dich von der KI interviewen", "Bonus · Säule 4 · Lektion 4.3")
    body(doc, "Statt vor einem leeren Blatt zu sitzen: lass dir deinen Business-Brief im Gespräch erarbeiten. Funktioniert in jedem Tool — sogar in der Gratis-Chatfunktion. Am Ende hast du EIN sauberes Dokument, das du überall als Wissen hinterlegst.")
    body(doc, "🔑 Mein Trick (genau so mach ich's selbst): Sag der KI, sie soll tun, als würde sie ein Buch über dich und dein Business schreiben — und so lange und vertieft fragen, bis sie wirklich ALLE Infos hat. Erst wenn sie dich so gut kennt wie eine Autorin ihre Hauptfigur, klingt sie später wie DU.", italic=True)
    h2(doc, "So gehst du vor")
    step(doc, "Öffne ein Chatfenster (Claude oder ChatGPT — Gratis reicht) oder dein Cowork.")
    step(doc, "Kopier den Prompt unten rein.")
    step(doc, "Beantworte die Fragen, eine nach der anderen — gib dir Zeit, antworte ehrlich und ausführlich.")
    step(doc, "Lass die KI ruhig nachbohren. Je tiefer sie fragt, desto besser wird dein Brief.")
    step(doc, "Am Schluss spuckt dir die KI „Mein Business-Brief\" als fertiges Dokument aus → kopieren + speichern.")
    h2(doc, "👇 Das hier kopierst du ins Chatfenster")
    codebox(doc, BRIEF)
    footer(doc)
    path = os.path.join(OUT, "Business-Brief-erarbeiten.docx")
    doc.save(path); print("saved", path)

# ===================== DOC 7: Cowork einrichten =====================
def build_cowork():
    doc = Document()
    h1(doc, "🖥 Claude Cowork einrichten — einmalig, ~10 Min", "Bonus · Säule 4 · Lektion 4.4")
    body(doc, "Das machst du EINMAL in Lektion 4.4. Danach laufen alle deine Bots darin — du musst es nie wieder einrichten. Kein Code, nur Knöpfe.")
    h2(doc, "Was du brauchst")
    bullet(doc, "Einen Computer (Mac oder Windows)")
    bullet(doc, "Ein Claude-Pro-Abo (~20-23 CHF/Monat) — die Gratis-Version reicht für Cowork nicht")
    bullet(doc, "Etwa 10 Minuten")
    h2(doc, "Schritt für Schritt")
    step(doc, "Claude Desktop installieren: Gehe auf claude.ai/download und lade die Desktop-App für dein System. Installieren, öffnen, mit deinem Claude-Konto anmelden.")
    step(doc, "Cowork öffnen: In der App findest du den Cowork-Bereich. (Falls du ihn nicht siehst: prüfe, ob dein Pro-Abo aktiv ist.)")
    step(doc, "Notion verbinden (der „Connector\" / Stecker): Einstellungen → Connectors → Notion auswählen → du wirst zu Notion geleitet → einloggen → freigeben, welche Seiten Claude sehen darf (mindestens dein Business-Brain + dein Privat-Bereich „Privat & Familie\"). Ab jetzt kann Claude in deinem Notion lesen.")
    step(doc, "Testen: Neuen Bot/Auftrag anlegen, frag z.B. „Was steht in meiner Notion-Wochenplanung?\". Wenn er antwortet → Verbindung läuft. 🎉")
    h2(doc, "Wichtig")
    bullet(doc, "Einmal eingerichtet, gilt für ALLE Bots (Cockpit, Haushalts-Helfer, Koch). Du verbindest Notion nicht jedes Mal neu.")
    bullet(doc, "Du gibst selbst frei, welche Seiten Claude sehen darf — er sieht nicht „dein ganzes Notion\", nur was du freischaltest.")
    bullet(doc, "Deinen Business-Brief legst du einmal in die Bot-Anweisung (oder als Notion-Seite, die der Bot liest) — auch nicht jedes Mal neu.")
    footer(doc)
    path = os.path.join(OUT, "Cowork-einrichten-Anleitung.docx")
    doc.save(path); print("saved", path)

for fn in (build_cockpit, build_claudecode, build_haushalt, build_koch, build_gratis, build_brief, build_cowork, build_notion_haushalt):
    try:
        fn()
    except PermissionError as e:
        print("SKIP (Datei offen in Word?):", e.filename)
print("ALL DONE")
